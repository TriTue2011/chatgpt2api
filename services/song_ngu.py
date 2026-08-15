"""Bản dịch SONG NGỮ đóng thành .docx — bản gốc và bản dịch nằm cạnh nhau.

Vì sao cần: bản dịch dài gửi thẳng vào khung chat thì Zalo cắt thành hàng chục
tin (giới hạn 2000 ký tự/tin), đọc rối và không đối chiếu được với bản gốc. Còn
gửi .txt chỉ có bản dịch thì mất luôn đường tra ngược khi máy dịch hiểu sai một
câu — mà máy dịch thì lúc nào cũng có câu hiểu sai.

Bố cục chọn kiểu XẾP CHỒNG (gốc rồi tới dịch, từng cặp) chứ không phải hai cột:
đọc trên điện thoại thì hai cột co lại còn vài chữ mỗi dòng. Bản gốc để chữ nhỏ
màu xám nên mắt lướt qua được khi không cần đối chiếu.

Ngắn thì KHÔNG đóng tệp: mở tệp để đọc ba câu là phiền hơn có ích.
"""

from __future__ import annotations

import io
import logging
from typing import Any, Iterable

logger = logging.getLogger(__name__)

#: Dài hơn ngần này ký tự thì đóng tệp thay vì nhắn thẳng. Lấy theo giới hạn
#: một tin Zalo (2000) — dưới mức đó vẫn gọn trong một tin, trên mức đó là bắt
#: đầu bị cắt vụn.
NGUONG_DONG_TEP = 1800

TEN_TIENG = {"vi": "Việt", "en": "Anh", "ja": "Nhật", "zh": "Trung",
             "ko": "Hàn", "auto": "tự nhận"}


def nen_dong_tep(chu: str) -> bool:
    """Bản dịch này nên đóng thành tệp hay nhắn thẳng?"""
    return len(str(chu or "")) > NGUONG_DONG_TEP


def _ten_tieng(ma: str) -> str:
    ma = str(ma or "").strip().lower()
    return TEN_TIENG.get(ma, ma or "?")


def docx_song_ngu(cap: Iterable[tuple[str, str]], *, nguon: str = "",
                  dich: str = "", tieu_de: str = "") -> bytes:
    """Các cặp (câu gốc, câu dịch) → tệp .docx trong bộ nhớ.

    Cặp nào thiếu một vế thì vẫn ghi vế còn lại — mất một câu vì dữ liệu lệch
    còn tệ hơn một dòng trống.
    """
    import docx
    from docx.shared import Pt, RGBColor

    d = docx.Document()
    d.add_heading(tieu_de or "Bản dịch song ngữ", level=1)
    dau = f"Tiếng {_ten_tieng(nguon)} → tiếng {_ten_tieng(dich)}"
    ghi = d.add_paragraph(dau)
    ghi.runs[0].italic = True

    for goc, ban in cap:
        goc, ban = str(goc or "").strip(), str(ban or "").strip()
        if not goc and not ban:
            continue
        if goc:
            p = d.add_paragraph()
            r = p.add_run(goc)
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        if ban:
            d.add_paragraph(ban)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def docx_mot_ban(chu: str, *, tieu_de: str = "", ghi_chu: str = "") -> bytes:
    """Một khối chữ → .docx thường (dùng cho bản chép lời, không dịch).

    Tách khỏi ``docx_song_ngu`` vì đây là tài liệu MỘT cột: chép lời xong người
    ta dán thẳng vào báo cáo, không cần cột đối chiếu nào.
    """
    import docx

    d = docx.Document()
    d.add_heading(tieu_de or "Bản chép lời", level=1)
    if ghi_chu:
        p = d.add_paragraph(ghi_chu)
        p.runs[0].italic = True
    for dong in str(chu or "").splitlines():
        if dong.strip():
            d.add_paragraph(dong.strip())
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def tach_cap(goc: str, ban: str) -> list[tuple[str, str]]:
    """Hai khối chữ → các cặp theo DÒNG.

    Dùng cho đường dịch đoạn chữ, nơi máy dịch trả về một khối. Số dòng hai bên
    lệch nhau thì ghép tới đâu hay tới đó rồi đổ phần thừa xuống cuối — thà lệch
    vài dòng cuối còn hơn cắt mất chữ.
    """
    a = [x for x in str(goc or "").splitlines() if x.strip()]
    b = [x for x in str(ban or "").splitlines() if x.strip()]
    ra = list(zip(a, b))
    if len(a) > len(b):
        ra += [(x, "") for x in a[len(b):]]
    elif len(b) > len(a):
        ra += [("", x) for x in b[len(a):]]
    return ra


def dong_goi(goc: str, ban: str, *, nguon: str = "", dich: str = "",
             tieu_de: str = "") -> dict[str, Any]:
    """Quyết định gửi thẳng hay đóng tệp, rồi trả thứ cần gửi.

    ``{"tep": bytes, "ten": str}`` nếu nên đóng tệp; ``{"chu": str}`` nếu ngắn.
    Đóng tệp lỗi (thiếu python-docx…) thì rơi về gửi chữ chứ không mất bản dịch.
    """
    if not nen_dong_tep(ban):
        return {"chu": ban}
    try:
        tep = docx_song_ngu(tach_cap(goc, ban), nguon=nguon, dich=dich,
                            tieu_de=tieu_de)
    except Exception as exc:
        logger.warning("đóng docx song ngữ lỗi: %s", str(exc)[:150])
        return {"chu": ban}
    return {"tep": tep, "ten": f"song-ngu.{dich or 'dich'}.docx"}
