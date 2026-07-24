"""Test pdf_intent: token gate (báo chi phí OCR trước khi đốt lượt vision)
và marker ảnh image:// (pdf_images.markdown_section)."""
import unittest

from services import pdf_intent as pi
from services import pdf_images


class CostNoteTests(unittest.TestCase):
    def test_no_info_no_note(self) -> None:
        msg = pi.ask_text("a.pdf", {"rag", "word"})
        self.assertNotIn("⚠️", msg)

    def test_digital_pdf_no_note(self) -> None:
        msg = pi.ask_text("a.pdf", {"rag", "word"},
                          {"pages": 30, "scanned": False, "ocr": False})
        self.assertNotIn("⚠️", msg)

    def test_small_scan_no_note(self) -> None:
        msg = pi.ask_text("a.pdf", {"rag", "word"},
                          {"pages": 2, "scanned": True, "ocr": True})
        self.assertNotIn("⚠️", msg)

    def test_big_scan_warns_pages_and_calls(self) -> None:
        msg = pi.ask_text("a.pdf", {"rag", "word"},
                          {"pages": 38, "scanned": True, "ocr": True})
        self.assertIn("⚠️", msg)
        self.assertIn("38 trang", msg)
        self.assertIn("lượt gọi", msg)
        # Vẫn giữ nguyên câu hỏi 1/2 phía trên.
        self.assertIn("1️⃣", msg)

    def test_over_cap_mentions_first_pages_only(self) -> None:
        msg = pi.ask_text("a.pdf", {"rag"},
                          {"pages": 120, "scanned": True, "ocr": True})
        self.assertIn("trang đầu", msg)


class ParseIntentTests(unittest.TestCase):
    def test_basic(self) -> None:
        # Số map theo INTENT_ORDER (rag_knowledge, rag_teacher, word, excel).
        self.assertEqual(pi.parse_intent("1"), "rag_knowledge")
        self.assertEqual(pi.parse_intent("2"), "rag_teacher")
        self.assertEqual(pi.parse_intent("3"), "word")
        self.assertIsNone(pi.parse_intent("xin chào"))


class ImageMarkerTests(unittest.TestCase):
    def test_markdown_section_format(self) -> None:
        sec = pdf_images.markdown_section(
            [{"id": "ab12", "page": 3, "caption": "Sơ đồ hệ thống điện"}])
        self.assertIn("## Hình ảnh trong tài liệu", sec)
        self.assertIn("![Sơ đồ hệ thống điện](image://ab12)", sec)
        self.assertIn("Trang 3", sec)

    def test_markdown_section_empty(self) -> None:
        self.assertEqual(pdf_images.markdown_section([]), "")

    def test_image_path_rejects_traversal(self) -> None:
        # UUID chỉ còn hex sau khi lọc → "../../etc" không trỏ ra ngoài thư mục.
        self.assertIsNone(pdf_images.image_path("../../etc/passwd"))


class MarkerScanTests(unittest.TestCase):
    def test_find_markers_and_dedup(self) -> None:
        t = ("Xem ![Sơ đồ điện](image://abcdef123456) và "
             "![Sơ đồ điện](image://abcdef123456) rồi ![](image://feedbeef99)")
        self.assertEqual(pdf_images.find_markers(t),
                         [("Sơ đồ điện", "abcdef123456"), ("", "feedbeef99")])

    def test_humanize_markers(self) -> None:
        t = "Trang 3 có ![Sơ đồ điện](image://abcdef123456)."
        out = pdf_images.humanize_markers(t)
        self.assertIn("[Hình: Sơ đồ điện]", out)
        self.assertNotIn("image://", out)

    def test_humanize_empty_caption(self) -> None:
        out = pdf_images.humanize_markers("![](image://abcdef123456)")
        self.assertEqual(out, "[Hình: trong tài liệu]")


class DigitalMarkdownTests(unittest.TestCase):
    """digital_pdf_markdown (PyMuPDF thuần): heading theo cỡ chữ + giữ thân bài."""

    def test_headings_and_body(self) -> None:
        import os
        import tempfile

        import fitz

        from services.pdf_to_word import digital_pdf_markdown

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 80), "BAO CAO TONG KET", fontsize=20)
        page.insert_text((72, 120), "Muc 1 gioi thieu", fontsize=14)
        for i in range(6):   # thân bài 11pt phải chiếm đa số ký tự
            page.insert_text((72, 160 + i * 16),
                             "Noi dung doan van binh thuong dong so " + str(i), fontsize=11)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        tmp.close()
        doc.save(tmp.name)
        doc.close()
        try:
            md = digital_pdf_markdown(tmp.name)
        finally:
            os.unlink(tmp.name)
        self.assertIn("# BAO CAO TONG KET", md)
        self.assertIn("## Muc 1 gioi thieu", md)
        self.assertIn("Noi dung doan van binh thuong dong so 0", md)

    def test_empty_pdf_returns_empty(self) -> None:
        import os
        import tempfile

        import fitz

        from services.pdf_to_word import digital_pdf_markdown

        doc = fitz.open()
        doc.new_page()
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        tmp.close()
        doc.save(tmp.name)
        doc.close()
        try:
            self.assertEqual(digital_pdf_markdown(tmp.name), "")
        finally:
            os.unlink(tmp.name)

    def test_list_bullets_normalized(self) -> None:
        import os
        import tempfile

        import fitz

        from services.pdf_to_word import digital_pdf_markdown

        doc = fitz.open()
        page = doc.new_page()
        # Thân bài nhiều ký tự hơn bullet để body size = 11.
        for i in range(8):
            page.insert_text((72, 80 + i * 14),
                             "Doan van than bai binh thuong so " + str(i), fontsize=11)
        page.insert_text((72, 220), "• Hang muc mot trong danh sach", fontsize=11)
        page.insert_text((72, 236), "• Hang muc hai trong danh sach", fontsize=11)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        tmp.close()
        doc.save(tmp.name)
        doc.close()
        try:
            md = digital_pdf_markdown(tmp.name)
        finally:
            os.unlink(tmp.name)
        self.assertIn("- Hang muc mot trong danh sach", md)
        self.assertIn("- Hang muc hai trong danh sach", md)


class AnalyzePdfHelpersTests(unittest.TestCase):
    def test_sample_page_indices_small(self) -> None:
        from services.pdf_to_word import _sample_page_indices
        self.assertEqual(_sample_page_indices(5), [0, 1, 2, 3, 4])

    def test_sample_page_indices_large_bounded(self) -> None:
        from services.pdf_to_word import _sample_page_indices
        idxs = _sample_page_indices(200, cap=12)
        self.assertLessEqual(len(idxs), 12)
        self.assertEqual(idxs[0], 0)
        self.assertIn(199, idxs)

    def test_text_looks_clean_latin(self) -> None:
        from services.pdf_to_word import _text_looks_clean
        clean = ("This is a clean English OCR layer with normal words and spaces. "
                 "It should pass the heuristic without Vietnamese diacritics.")
        self.assertTrue(_text_looks_clean(clean))

    def test_text_looks_clean_rejects_mojibake(self) -> None:
        from services.pdf_to_word import _text_looks_clean
        bad = "a b c d e f g h i j k l m n o p q r s t u v w x y z " * 2
        self.assertFalse(_text_looks_clean(bad))
        self.assertFalse(_text_looks_clean("short"))

    def test_analyze_digital_pdf(self) -> None:
        import os
        import tempfile

        import fitz

        from services.pdf_to_word import analyze_pdf

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello digital PDF content for analysis test.", fontsize=12)
        page.insert_text((72, 100), "Second line of body text here too.", fontsize=12)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        tmp.close()
        doc.save(tmp.name)
        doc.close()
        try:
            info = analyze_pdf(tmp.name)
        finally:
            os.unlink(tmp.name)
        self.assertFalse(info["scanned"])
        self.assertEqual(info["text_quality"], "good")
        self.assertEqual(info["pages"], 1)


class MdIntoDocxTests(unittest.TestCase):
    def test_lists_and_bold_to_docx(self) -> None:
        import os
        import tempfile

        from docx import Document

        from services.pdf_to_word import _markdown_to_docx

        md = "# Title\n\nBody with **bold** word.\n\n- item one\n- item two\n\n1. first\n"
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        tmp.close()
        try:
            _markdown_to_docx(md, tmp.name)
            doc = Document(tmp.name)
            texts = [p.text for p in doc.paragraphs if p.text.strip()]
            joined = "\n".join(texts)
            self.assertIn("Title", joined)
            self.assertIn("bold", joined)
            self.assertIn("item one", joined)
            # Có ít nhất một run bold.
            has_bold = any(r.bold for p in doc.paragraphs for r in p.runs)
            self.assertTrue(has_bold)
        finally:
            os.unlink(tmp.name)


if __name__ == "__main__":
    unittest.main()
